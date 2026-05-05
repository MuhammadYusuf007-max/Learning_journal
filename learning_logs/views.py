import os
import io
import json
import re
from dotenv import load_dotenv
from openai import OpenAI

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Count
from django.http import Http404, FileResponse, HttpResponse
from django.template.defaultfilters import striptags
from django.utils import timezone

from django.http import StreamingHttpResponse
from django.views.decorators.http import require_POST

from .models import Topic, Entry, QuizAttempt, Flashcard, QAExchange, Tag, AIUsage
from .forms import TopicForm, EntryForm, FlashcardForm

# Libraries for file export
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from docx import Document

# --- AI SETUP ---
load_dotenv()

_ai_client = None


def _get_ai_client():
    """Lazily build the AI client so the app can boot without an API key."""
    global _ai_client
    if _ai_client is not None:
        return _ai_client

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return None

    _ai_client = OpenAI(
        api_key=api_key,
        base_url="https://api.groq.com/openai/v1",
    )
    return _ai_client


def _log_usage(user, feature, response):
    """Record token usage for one AI call. Best-effort, swallows all errors."""
    try:
        usage = getattr(response, 'usage', None)
        if usage is None:
            return
        AIUsage.objects.create(
            user=user if (user and user.is_authenticated) else None,
            feature=feature or 'other',
            prompt_tokens=getattr(usage, 'prompt_tokens', 0) or 0,
            completion_tokens=getattr(usage, 'completion_tokens', 0) or 0,
            total_tokens=getattr(usage, 'total_tokens', 0) or 0,
        )
    except Exception as e:
        print(f"AI usage logging error: {e}")


def _call_ai(system_msg, user_msg, user=None, feature=None):
    """Low-level helper that makes one chat completion call."""
    client = _get_ai_client()
    if client is None:
        return "AI features disabled: GROQ_API_KEY is not set."
    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
        )
        _log_usage(user, feature, response)
        return response.choices[0].message.content
    except Exception as e:
        print(f"AI Error: {e}")
        return "Content unavailable at this time."


def generate_ai_content(text, mode="summary", user=None):
    """
    Refined AI helper to handle different tasks.
    Modes: 'summary', 'master', 'quiz'
    """
    if mode == "quiz":
        system_msg = "You are a demanding university professor."
        user_msg = f"Based on the following notes, generate 5 challenging multiple-choice or open-ended study questions. Do not provide answers, just the questions. Notes: {text}"
        feature = 'quiz'
    elif mode == "master":
        system_msg = "You are a helpful academic assistant."
        user_msg = f"Provide a cohesive 3-paragraph summary of this learning progress: {text}"
        feature = 'master'
    else:
        system_msg = "You are a helpful assistant."
        user_msg = f"Summarize this in one short sentence: {text}"
        feature = 'summary'

    return _call_ai(system_msg, user_msg, user=user, feature=feature)


def generate_ai_quiz_json(notes_text, num_questions=5, user=None):
    """
    Ask the AI for a structured multiple-choice quiz in JSON form.
    Returns a list of question dicts, or an empty list on failure.
    """
    MAX_NOTES_CHARS = 12000
    if len(notes_text) > MAX_NOTES_CHARS:
        notes_text = notes_text[:MAX_NOTES_CHARS]

    system_msg = (
        "You are a strict but fair tutor that creates multiple-choice quizzes. "
        "You MUST respond with ONLY a valid JSON array, no prose, no markdown fences. "
        "Each item must have exactly these keys: question (string), options (array of "
        "exactly 4 strings), correct_index (integer 0-3), explanation (short string)."
    )
    user_msg = (
        f"Create exactly {num_questions} multiple-choice questions based ONLY on the "
        f"notes below. Each question must have 4 options. Indicate the correct answer "
        f"by its index (0-3). Add a one-sentence explanation.\n\n"
        f"NOTES:\n{notes_text}\n\n"
        f"Return JSON only."
    )

    raw = _call_ai(system_msg, user_msg, user=user, feature='quiz')
    if not raw or raw.startswith("AI features disabled"):
        return []

    # Strip optional ```json fences the model may add despite instructions.
    cleaned = raw.strip()
    fence_match = re.search(r"```(?:json)?\s*(.*?)```", cleaned, re.DOTALL)
    if fence_match:
        cleaned = fence_match.group(1).strip()

    # If the model added prose before/after, try to extract the first JSON array.
    if not cleaned.startswith('['):
        bracket_match = re.search(r"\[\s*\{.*\}\s*\]", cleaned, re.DOTALL)
        if bracket_match:
            cleaned = bracket_match.group(0)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as e:
        print(f"Quiz JSON parse error: {e}\nRaw: {raw[:500]}")
        return []

    # Validate shape and keep only well-formed items.
    valid = []
    for item in data if isinstance(data, list) else []:
        if not isinstance(item, dict):
            continue
        q = item.get('question')
        opts = item.get('options')
        ci = item.get('correct_index')
        expl = item.get('explanation', '')
        if (
            isinstance(q, str) and q.strip()
            and isinstance(opts, list) and len(opts) == 4
            and all(isinstance(o, str) for o in opts)
            and isinstance(ci, int) and 0 <= ci <= 3
        ):
            valid.append({
                'question': q.strip(),
                'options': [o.strip() for o in opts],
                'correct_index': ci,
                'explanation': str(expl).strip(),
            })
    return valid


_TAG_SAFE_RE = re.compile(r'[^a-z0-9\-]+')


def _normalize_tag(name):
    """Lowercase, slugify, strip to a safe tag name. Returns '' if invalid."""
    if not isinstance(name, str):
        return ''
    s = name.strip().lower().replace(' ', '-')
    s = _TAG_SAFE_RE.sub('', s)
    s = s.strip('-')
    return s[:50]


def generate_ai_tags_json(text, num_tags=4, user=None):
    """Ask the AI for short tag strings as a JSON array. Returns list of normalized tag names."""
    MAX_CHARS = 6000
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]

    system_msg = (
        "You generate short, lowercase, single-word or hyphenated tags for study notes. "
        "Respond with ONLY a JSON array of strings, no prose, no markdown. "
        "Each tag must be 1-3 words, lowercase, hyphen-separated, max 30 chars. "
        "Examples: \"machine-learning\", \"calculus\", \"history\", \"databases\"."
    )
    user_msg = (
        f"Generate exactly {num_tags} concise tags for the following note:\n\n"
        f"{text}\n\n"
        f"Return JSON only."
    )

    raw = _call_ai(system_msg, user_msg, user=user, feature='tags')
    if not raw or raw.startswith("AI features disabled"):
        return []

    cleaned = raw.strip()
    fence_match = re.search(r"```(?:json)?\s*(.*?)```", cleaned, re.DOTALL)
    if fence_match:
        cleaned = fence_match.group(1).strip()
    if not cleaned.startswith('['):
        bracket_match = re.search(r"\[.*\]", cleaned, re.DOTALL)
        if bracket_match:
            cleaned = bracket_match.group(0)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        return []

    if not isinstance(data, list):
        return []

    seen = set()
    result = []
    for item in data:
        norm = _normalize_tag(item)
        if norm and norm not in seen:
            seen.add(norm)
            result.append(norm)
    return result[:num_tags]


def _attach_ai_tags(entry, user):
    """Generate tags for an entry's text and link them. Best-effort, no exceptions raised."""
    try:
        names = generate_ai_tags_json(striptags(entry.text), num_tags=4, user=user)
        for name in names:
            tag, _ = Tag.objects.get_or_create(name=name, owner=user)
            entry.tags.add(tag)
        return names
    except Exception as e:
        print(f"Auto-tag error: {e}")
        return []


def generate_ai_flashcards_json(notes_text, num_cards=8, user=None):
    """
    Ask the AI for flashcards in JSON form.
    Returns a list of {"front": str, "back": str} dicts, or [] on failure.
    """
    MAX_NOTES_CHARS = 12000
    if len(notes_text) > MAX_NOTES_CHARS:
        notes_text = notes_text[:MAX_NOTES_CHARS]

    system_msg = (
        "You are a study coach that creates concise flashcards. "
        "Respond with ONLY a valid JSON array, no prose, no markdown fences. "
        "Each item must have exactly: front (a short question or prompt, max 200 "
        "chars) and back (a clear answer, max 500 chars)."
    )
    user_msg = (
        f"Create exactly {num_cards} flashcards based ONLY on the notes below. "
        f"Cover the most important concepts. Make 'front' a question or term, "
        f"and 'back' the answer or definition.\n\n"
        f"NOTES:\n{notes_text}\n\n"
        f"Return JSON only."
    )

    raw = _call_ai(system_msg, user_msg, user=user, feature='flashcards')
    if not raw or raw.startswith("AI features disabled"):
        return []

    cleaned = raw.strip()
    fence_match = re.search(r"```(?:json)?\s*(.*?)```", cleaned, re.DOTALL)
    if fence_match:
        cleaned = fence_match.group(1).strip()
    if not cleaned.startswith('['):
        bracket_match = re.search(r"\[\s*\{.*\}\s*\]", cleaned, re.DOTALL)
        if bracket_match:
            cleaned = bracket_match.group(0)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as e:
        print(f"Flashcard JSON parse error: {e}\nRaw: {raw[:500]}")
        return []

    valid = []
    for item in data if isinstance(data, list) else []:
        if not isinstance(item, dict):
            continue
        front = item.get('front')
        back = item.get('back')
        if isinstance(front, str) and front.strip() and isinstance(back, str) and back.strip():
            valid.append({
                'front': front.strip()[:300],
                'back': back.strip()[:2000],
            })
    return valid


def _build_qa_messages(context_text, question):
    """Build the system + user messages for a Q&A request (shared by sync + streaming)."""
    MAX_CONTEXT_CHARS = 12000
    if len(context_text) > MAX_CONTEXT_CHARS:
        context_text = context_text[:MAX_CONTEXT_CHARS] + "\n[...notes truncated...]"

    system_msg = (
        "You are a helpful tutor that answers questions strictly using the "
        "student's own notes provided below. If the answer cannot be found in "
        "the notes, say so honestly and briefly suggest what they could study "
        "next. Be concise, accurate, and use bullet points when useful."
    )
    user_msg = (
        f"NOTES:\n{context_text}\n\n"
        f"QUESTION:\n{question}\n\n"
        "Answer using only the notes above."
    )
    return system_msg, user_msg


def generate_ai_qa(context_text, question):
    """
    Answer a user's question using the topic's notes as context (basic RAG).
    Blocking version (returns the full answer string).
    """
    system_msg, user_msg = _build_qa_messages(context_text, question)
    return _call_ai(system_msg, user_msg)


def stream_ai_qa(context_text, question, user=None):
    """
    Streaming version of generate_ai_qa. Yields text chunks as they arrive
    from the model so the UI can render the answer progressively.
    """
    client = _get_ai_client()
    if client is None:
        yield "AI features disabled: GROQ_API_KEY is not set."
        return

    system_msg, user_msg = _build_qa_messages(context_text, question)

    try:
        stream = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            stream=True,
            stream_options={"include_usage": True},
        )
        usage = None
        for chunk in stream:
            try:
                if chunk.choices:
                    delta = chunk.choices[0].delta.content
                else:
                    delta = None
            except (AttributeError, IndexError):
                delta = None
            if delta:
                yield delta
            # The final chunk in a stream-with-usage call has empty choices and a populated `usage`.
            if getattr(chunk, 'usage', None):
                usage = chunk.usage

        if usage is not None:
            try:
                AIUsage.objects.create(
                    user=user if (user and user.is_authenticated) else None,
                    feature='qa',
                    prompt_tokens=getattr(usage, 'prompt_tokens', 0) or 0,
                    completion_tokens=getattr(usage, 'completion_tokens', 0) or 0,
                    total_tokens=getattr(usage, 'total_tokens', 0) or 0,
                )
            except Exception as e:
                print(f"AI streaming usage logging error: {e}")
    except Exception as e:
        print(f"AI streaming error: {e}")
        yield f"\n\n[Error: AI service is unavailable right now.]"

# --- VIEWS ---

def index(request):
    """The home page for Learning Log"""
    return render(request, 'learning_logs/index.html')

@login_required
def topics(request):
    """Show all topics"""
    topics = Topic.objects.filter(owner=request.user).order_by('date_added')
    context = {'topics': topics}
    return render(request, 'learning_logs/topics.html', context)

@login_required
def topic(request, topic_id):
    """Show a single topic and a paginated list of its entries."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    all_entries = topic.entry_set.order_by('-date_added')

    paginator = Paginator(all_entries, 10)  # 10 entries per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'topic': topic,
        'entries': page_obj,
        'page_obj': page_obj,
        'is_paginated': page_obj.has_other_pages(),
    }
    return render(request, 'learning_logs/topic.html', context)

@login_required
def topic_summary(request, topic_id):
    """Generates a master summary and handles PDF/Docx exports."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    entries = topic.entry_set.all().order_by('date_added')
    
    # Combine entries for the AI
    raw_text = " ".join([striptags(e.text) for e in entries])
    summary_text = generate_ai_content(raw_text, mode="master", user=request.user)

    export_format = request.GET.get('export')

    # Handle Word Export
    if export_format == 'docx':
        doc = Document()
        doc.add_heading(f'Learning Summary: {topic.text}', 0)
        doc.add_paragraph(summary_text)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename=f'{topic.text}_summary.docx')

    # Handle PDF Export
    elif export_format == 'pdf':
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        p.setFont("Helvetica-Bold", 16)
        p.drawString(72, height - 72, f"Learning Summary: {topic.text}")
        
        p.setFont("Helvetica", 12)
        text_object = p.beginText(72, height - 100)
        
        # Wrapping text so it doesn't go off the page
        lines = simpleSplit(summary_text, "Helvetica", 12, width - 144)
        for line in lines:
            if text_object.getY() < 72: # Create new page if full
                p.drawText(text_object)
                p.showPage()
                p.setFont("Helvetica", 12)
                text_object = p.beginText(72, height - 72)
            text_object.textLine(line)
            
        p.drawText(text_object)
        p.showPage()
        p.save()
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename=f'{topic.text}_summary.pdf')

    return render(request, 'learning_logs/topic_summary.html', {
        'topic': topic,
        'summary': summary_text
    })

@login_required
def new_topic(request):
    """Add a new topic."""
    if request.method != 'POST':
        form = TopicForm()
    else:
        form = TopicForm(data=request.POST)
        if form.is_valid():
            new_topic = form.save(commit=False)
            new_topic.owner = request.user
            new_topic.save()
            messages.success(request, f'Topic "{new_topic.text}" created.')
            return redirect('learning_logs:topics')
    context = {'form': form}
    return render(request, 'learning_logs/new_topic.html', context)

@login_required
def new_entry(request, topic_id):
    """Add a new entry for a topic."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    if request.method != 'POST':
        form = EntryForm()
    else:
        form = EntryForm(data=request.POST)
        if form.is_valid():
            new_entry = form.save(commit=False)
            new_entry.topic = topic
            clean_text = striptags(new_entry.text)
            new_entry.ai_summary = generate_ai_content(clean_text, mode="summary", user=request.user)
            new_entry.save()
            tag_names = _attach_ai_tags(new_entry, request.user)
            if tag_names:
                messages.success(request, f'Entry added. Auto-tagged with: {", ".join(tag_names)}.')
            else:
                messages.success(request, 'Entry added.')
            return redirect('learning_logs:topic', topic_id=topic_id)
    context = {'topic': topic, 'form': form}
    return render(request, 'learning_logs/new_entry.html', context)

@login_required
def edit_entry(request, entry_id):
    """Edit an existing entry, including its tags."""
    entry = get_object_or_404(Entry, id=entry_id)
    topic = entry.topic
    if topic.owner != request.user:
        raise Http404

    if request.method != 'POST':
        form = EntryForm(instance=entry)
    else:
        form = EntryForm(instance=entry, data=request.POST)
        if form.is_valid():
            edited_entry = form.save(commit=False)
            clean_text = striptags(edited_entry.text)
            edited_entry.ai_summary = generate_ai_content(clean_text, mode="summary", user=request.user)
            edited_entry.save()

            # Update tags from a comma-separated text field.
            raw_tags = request.POST.get('tags_text', '')
            new_tag_names = []
            for piece in raw_tags.split(','):
                norm = _normalize_tag(piece)
                if norm:
                    new_tag_names.append(norm)

            edited_entry.tags.clear()
            for name in new_tag_names:
                tag, _ = Tag.objects.get_or_create(name=name, owner=request.user)
                edited_entry.tags.add(tag)

            messages.success(request, 'Entry updated.')
            return redirect('learning_logs:topic', topic_id=topic.id)

    context = {
        'entry': entry,
        'topic': topic,
        'form': form,
        'tags_text': ', '.join(entry.tags.values_list('name', flat=True)),
    }
    return render(request, 'learning_logs/edit_entry.html', context)

@login_required
def edit_topic(request, topic_id):
    """Edit an existing topic."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    if request.method != 'POST':
        form = TopicForm(instance=topic)
    else:
        form = TopicForm(instance=topic, data=request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Topic updated.')
            return redirect('learning_logs:topic', topic_id=topic.id)
    context = {'topic': topic, 'form': form}
    return render(request, 'learning_logs/edit_topic.html', context)

@login_required
def delete_topic(request, topic_id):
    """Delete an existing topic."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    if request.method == 'POST':
        topic_name = topic.text
        topic.delete()
        messages.success(request, f'Topic "{topic_name}" deleted.')
        return redirect('learning_logs:topics')
    return render(request, 'learning_logs/delete_topic.html', {'topic': topic})

@login_required
def topic_flashcards(request, topic_id):
    """List a topic's flashcards with management actions."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    cards = topic.flashcards.all()
    return render(request, 'learning_logs/flashcards_list.html', {
        'topic': topic,
        'cards': cards,
    })


@login_required
def generate_flashcards(request, topic_id):
    """POST endpoint that uses the AI to generate flashcards for the topic."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)

    if request.method != 'POST':
        return redirect('learning_logs:topic_flashcards', topic_id=topic.id)

    entries = topic.entry_set.all()
    if not entries.exists():
        messages.warning(request, 'Add at least one entry before generating flashcards.')
        return redirect('learning_logs:topic_flashcards', topic_id=topic.id)

    raw_text = " ".join(striptags(e.text) for e in entries)
    cards_data = generate_ai_flashcards_json(raw_text, num_cards=8, user=request.user)

    if not cards_data:
        messages.error(request, "Couldn't generate flashcards right now. Please try again in a moment.")
        return redirect('learning_logs:topic_flashcards', topic_id=topic.id)

    Flashcard.objects.bulk_create([
        Flashcard(topic=topic, front=c['front'], back=c['back'])
        for c in cards_data
    ])
    messages.success(request, f'Generated {len(cards_data)} flashcards.')
    return redirect('learning_logs:topic_flashcards', topic_id=topic.id)


@login_required
def new_flashcard(request, topic_id):
    """Manually add a single flashcard."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)

    if request.method != 'POST':
        form = FlashcardForm()
    else:
        form = FlashcardForm(data=request.POST)
        if form.is_valid():
            card = form.save(commit=False)
            card.topic = topic
            card.save()
            messages.success(request, 'Flashcard added.')
            return redirect('learning_logs:topic_flashcards', topic_id=topic.id)

    return render(request, 'learning_logs/new_flashcard.html', {
        'topic': topic,
        'form': form,
    })


@login_required
def delete_flashcard(request, card_id):
    """Delete a single flashcard (must be owned by current user)."""
    card = get_object_or_404(Flashcard, id=card_id, topic__owner=request.user)
    topic_id = card.topic.id
    if request.method == 'POST':
        card.delete()
        messages.success(request, 'Flashcard deleted.')
    return redirect('learning_logs:topic_flashcards', topic_id=topic_id)


@login_required
def review_flashcards(request, topic_id):
    """One-card-at-a-time review session. Tracks 'got it' / 'study again'."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    cards = list(topic.flashcards.all())

    if not cards:
        messages.warning(request, 'Add or generate some flashcards first.')
        return redirect('learning_logs:topic_flashcards', topic_id=topic.id)

    session_key = f'flashcard_session_{topic_id}'

    if request.method == 'POST':
        action = request.POST.get('action')
        card_id = request.POST.get('card_id')
        if card_id and action in ('correct', 'incorrect'):
            try:
                card = Flashcard.objects.get(id=int(card_id), topic=topic)
                card.times_seen += 1
                if action == 'correct':
                    card.times_correct += 1
                card.save()
            except (Flashcard.DoesNotExist, ValueError):
                pass

        # Advance the index in the session.
        idx = request.session.get(session_key, 0) + 1
        request.session[session_key] = idx

        if idx >= len(cards):
            request.session[session_key] = 0
            return redirect('learning_logs:topic_flashcards', topic_id=topic.id)
        return redirect('learning_logs:review_flashcards', topic_id=topic.id)

    idx = request.session.get(session_key, 0)
    if idx >= len(cards):
        idx = 0
        request.session[session_key] = 0

    return render(request, 'learning_logs/review_flashcards.html', {
        'topic': topic,
        'card': cards[idx],
        'index': idx + 1,
        'total': len(cards),
    })


def _build_qa_context(topic):
    """Concatenate a topic's entries into a single context block for the AI."""
    entries = topic.entry_set.all().order_by('date_added')
    return "\n\n".join(
        f"[Entry {i+1}, {e.date_added:%Y-%m-%d}]\n{striptags(e.text)}"
        for i, e in enumerate(entries)
    )


@login_required
def topic_qa(request, topic_id):
    """Chat-style Q&A page (loads history from the DB; streaming happens via JS)."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    entries_exist = topic.entry_set.exists()
    entry_count = topic.entry_set.count()

    if request.method == 'POST' and request.POST.get('action') == 'clear':
        QAExchange.objects.filter(user=request.user, topic=topic).delete()
        messages.success(request, 'Conversation cleared.')
        return redirect('learning_logs:topic_qa', topic_id=topic.id)

    history = QAExchange.objects.filter(user=request.user, topic=topic).order_by('created_at')

    return render(request, 'learning_logs/topic_qa.html', {
        'topic': topic,
        'history': history,
        'has_entries': entries_exist,
        'entry_count': entry_count,
    })


@login_required
@require_POST
def topic_qa_stream(request, topic_id):
    """
    Stream an AI answer to a question for this topic.
    Returns a StreamingHttpResponse of plain-text chunks. The full answer is
    persisted to a QAExchange row when the stream finishes.
    """
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)

    question = request.POST.get('question', '').strip()
    if not question:
        return HttpResponse('Please enter a question.', status=400)

    if not topic.entry_set.exists():
        return HttpResponse(
            'This topic has no entries yet. Add some notes before asking questions.',
            status=400,
        )

    context_text = _build_qa_context(topic)
    exchange = QAExchange.objects.create(
        user=request.user,
        topic=topic,
        question=question,
    )

    user = request.user

    def generate():
        full_answer = []
        try:
            for chunk in stream_ai_qa(context_text, question, user=user):
                full_answer.append(chunk)
                yield chunk
        finally:
            exchange.answer = ''.join(full_answer)
            exchange.completed_at = timezone.now()
            exchange.save(update_fields=['answer', 'completed_at'])

    response = StreamingHttpResponse(generate(), content_type='text/plain; charset=utf-8')
    # Hint to nginx/reverse proxies not to buffer the response so chunks reach
    # the browser immediately.
    response['X-Accel-Buffering'] = 'no'
    response['Cache-Control'] = 'no-cache'
    return response


@login_required
def tags_list(request):
    """Show all of the user's tags with usage counts."""
    tags = (
        Tag.objects
        .filter(owner=request.user)
        .annotate(num_entries=Count('entries'))
        .order_by('-num_entries', 'name')
    )
    return render(request, 'learning_logs/tags_list.html', {'tags': tags})


@login_required
def entries_by_tag(request, tag_id):
    """Show all entries (across topics) tagged with `tag_id`, scoped to this user."""
    tag = get_object_or_404(Tag, id=tag_id, owner=request.user)
    entries = (
        tag.entries
        .filter(topic__owner=request.user)
        .select_related('topic')
        .order_by('-date_added')
    )
    return render(request, 'learning_logs/entries_by_tag.html', {
        'tag': tag,
        'entries': entries,
    })


@login_required
def search(request):
    """Search the current user's topics and entries."""
    query = request.GET.get('q', '').strip()

    topic_results = []
    entry_results = []

    if query:
        topic_results = (
            Topic.objects
            .filter(owner=request.user, text__icontains=query)
            .order_by('-date_added')
        )
        entry_results = (
            Entry.objects
            .filter(topic__owner=request.user, text__icontains=query)
            .select_related('topic')
            .order_by('-date_added')
        )

    context = {
        'query': query,
        'topic_results': topic_results,
        'entry_results': entry_results,
        'total': len(topic_results) + len(entry_results),
    }
    return render(request, 'learning_logs/search_results.html', context)


@login_required
def topic_quiz(request, topic_id):
    """Start a new quiz attempt: generate questions, save, then redirect to take it."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    entries = topic.entry_set.all()

    if not entries.exists():
        messages.warning(request, 'Add at least one entry before taking a quiz.')
        return redirect('learning_logs:topic', topic_id=topic.id)

    raw_text = " ".join(striptags(e.text) for e in entries)
    questions = generate_ai_quiz_json(raw_text, num_questions=5, user=request.user)

    if not questions:
        messages.error(request, "Couldn't generate a quiz right now. Please try again in a moment.")
        return redirect('learning_logs:topic', topic_id=topic.id)

    attempt = QuizAttempt.objects.create(
        user=request.user,
        topic=topic,
        questions_data=questions,
        total=len(questions),
    )
    return redirect('learning_logs:take_quiz', attempt_id=attempt.id)


@login_required
def take_quiz(request, attempt_id):
    """Display the quiz form and grade it on submit."""
    attempt = get_object_or_404(QuizAttempt, id=attempt_id, user=request.user)

    if attempt.completed:
        return redirect('learning_logs:quiz_result', attempt_id=attempt.id)

    if request.method == 'POST':
        answers = []
        score = 0
        for i, q in enumerate(attempt.questions_data):
            raw = request.POST.get(f'q_{i}')
            try:
                picked = int(raw) if raw is not None else None
            except (TypeError, ValueError):
                picked = None
            answers.append(picked)
            if picked == q.get('correct_index'):
                score += 1

        attempt.answers_data = answers
        attempt.score = score
        attempt.completed = True
        attempt.completed_at = timezone.now()
        attempt.save()
        messages.success(request, f'You scored {score} / {attempt.total} ({attempt.percentage}%).')
        return redirect('learning_logs:quiz_result', attempt_id=attempt.id)

    return render(request, 'learning_logs/take_quiz.html', {
        'attempt': attempt,
        'topic': attempt.topic,
    })


@login_required
def quiz_result(request, attempt_id):
    """Show the graded results for a completed quiz attempt."""
    attempt = get_object_or_404(QuizAttempt, id=attempt_id, user=request.user)

    if not attempt.completed:
        return redirect('learning_logs:take_quiz', attempt_id=attempt.id)

    # Build a list of {question, options, correct_index, picked, explanation, is_correct}
    review = []
    for i, q in enumerate(attempt.questions_data):
        picked = attempt.answers_data[i] if i < len(attempt.answers_data) else None
        review.append({
            'index': i,
            'question': q.get('question', ''),
            'options': q.get('options', []),
            'correct_index': q.get('correct_index'),
            'picked': picked,
            'explanation': q.get('explanation', ''),
            'is_correct': picked == q.get('correct_index'),
        })

    return render(request, 'learning_logs/quiz_result.html', {
        'attempt': attempt,
        'topic': attempt.topic,
        'review': review,
    })


@login_required
def quiz_history(request, topic_id):
    """List the user's past quiz attempts on a topic."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    attempts = QuizAttempt.objects.filter(user=request.user, topic=topic, completed=True)

    best = max((a.percentage for a in attempts), default=0)
    avg = round(sum(a.percentage for a in attempts) / len(attempts), 1) if attempts else 0

    return render(request, 'learning_logs/quiz_history.html', {
        'topic': topic,
        'attempts': attempts,
        'best': best,
        'avg': avg,
    })